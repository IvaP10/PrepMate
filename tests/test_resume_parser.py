import base64
import io
import tempfile
import unittest
import zipfile
from pathlib import Path
from unittest.mock import patch

from pre_interview import _extract_target_role, extract_resume_with_rules
from resume_parser import parse_resume_structured, validate_resume_bytes, _normalize_text


class ResumeParserTests(unittest.TestCase):
    def test_resume_container_signature_must_match_extension(self):
        with self.assertRaises(ValueError):
            validate_resume_bytes(b"not a pdf", ".pdf")
        with self.assertRaises(ValueError):
            validate_resume_bytes(b"%PDF-1.7 fake", ".docx")

    def test_docx_zip_bomb_ratio_is_rejected_before_python_docx(self):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            archive.writestr("[Content_Types].xml", "<Types />")
            archive.writestr("word/document.xml", "A" * 500_000)

        with self.assertRaisesRegex(ValueError, "compression ratio"):
            validate_resume_bytes(buffer.getvalue(), ".docx")

    def test_docx_archive_path_traversal_is_rejected(self):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("[Content_Types].xml", "<Types />")
            archive.writestr("word/document.xml", "<document />")
            archive.writestr("word/../../outside.xml", "<outside />")

        with self.assertRaisesRegex(ValueError, "unsafe file path"):
            validate_resume_bytes(buffer.getvalue(), ".docx")

    def test_docx_external_entity_declarations_are_rejected(self):
        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as archive:
            archive.writestr("[Content_Types].xml", "<Types />")
            archive.writestr(
                "word/document.xml",
                '<!DOCTYPE x [<!ENTITY ex SYSTEM "file:///etc/passwd">]><document>&ex;</document>',
            )

        with self.assertRaisesRegex(ValueError, "unsafe XML"):
            validate_resume_bytes(buffer.getvalue(), ".docx")

    def test_extract_target_role_from_experience_header(self):
        text = """
        Jane Doe
        EXPERIENCE
        Senior Software Engineer at Acme Corp
        Jan 2022 - Present
        Built APIs
        """
        sections = {
            "experience": ["Senior Software Engineer at Acme Corp", "Jan 2022 - Present", "Built APIs"],
            "summary": [],
        }
        role = _extract_target_role(text, sections)
        self.assertEqual(role, "Senior Software Engineer")

    def test_rule_extraction_includes_target_role(self):
        text = """
        Alex Kim
        alex@example.com
        EXPERIENCE
        Data Analyst at Contoso
        2021 - 2023
        SQL, Python, Tableau
        SKILLS
        Python, SQL, Tableau
        """
        profile = extract_resume_with_rules(
            text,
            {"email": "alex@example.com", "phone": None},
            {"linkedin": None, "github": None, "portfolio": None},
            {"parser": "pypdfium2", "links": []},
        )
        self.assertEqual(profile.get("target_role"), "Data Analyst")
        self.assertIn("Python", profile.get("skills", []))

    def test_parse_pdf_fast_skips_ocr_metadata(self):
        try:
            import pypdfium2
        except ImportError:
            self.skipTest("pypdfium2 not installed")

        # Small one-page fixture with selectable text. Keeping it inline avoids
        # requiring a PDF authoring library in the development or test install.
        pdf_bytes = base64.b64decode(
            "JVBERi0xLjcKJcK1wrYKJSBXcml0dGVuIGJ5IE11UERGIDEuMjcuMgoKMSAwIG9iago8PC9UeXBlL0NhdGFsb2cvUGFnZXMgMiAwIFIvSW5mbzw8L1Byb2R1Y2VyKE11UERGIDEuMjcuMik+Pj4+CmVuZG9iagoKMiAwIG9iago8PC9UeXBlL1BhZ2VzL0NvdW50IDEvS2lkc1s0IDAgUl0+PgplbmRvYmoKCjMgMCBvYmoKPDwvRm9udDw8L2hlbHYgNSAwIFI+Pj4+CmVuZG9iagoKNCAwIG9iago8PC9UeXBlL1BhZ2UvTWVkaWFCb3hbMCAwIDU5NSA4NDJdL1JvdGF0ZSAwL1Jlc291cmNlcyAzIDAgUi9QYXJlbnQgMiAwIFIvQ29udGVudHNbNiAwIFJdPj4KZW5kb2JqCgo1IDAgb2JqCjw8L1R5cGUvRm9udC9TdWJ0eXBlL1R5cGUxL0Jhc2VGb250L0hlbHZldGljYS9FbmNvZGluZy9XaW5BbnNpRW5jb2Rpbmc+PgplbmRvYmoKCjYgMCBvYmoKPDwvTGVuZ3RoIDE1My9GaWx0ZXIvRmxhdGVEZWNvZGU+PgpzdHJlYW0KeNoljTELwlAMhPf8isyCmry+XFooDqKLm/C24iDtKx10cPH3m1YCgdx9uaMPnQspS4yyJ3YXLm86LvX1ZVUuMw99fmJGi5rEGkzoPKM9PcqNhPdqB9XM5UJDH+4MeHaHeoIlyYYKRxfbYJ6ShJNX3VsoJheMG7cSzZZZdtFo4ltL5NU0Rm+kBd/Eb1xZze1PXwvd6QdOFiyzCmVuZHN0cmVhbQplbmRvYmoKCnhyZWYKMCA3CjAwMDAwMDAwMDAgNjU1MzUgZiAKMDAwMDAwMDA0MiAwMDAwMCBuIAowMDAwMDAxMjAgMDAwMDAgbiAKMDAwMDAwMTcyIDAwMDAwIG4gCjAwMDAwMDIxMyAwMDAwMCBuIAowMDAwMDAwMzIwIDAwMDAwIG4gCjAwMDAwMDQwOSAwMDAwMCBuIAAKCnRyYWlsZXIKPDwvU2l6ZSA3L1Jvb3QgMSAwIFIvSURbPDc5NkZDMjhENThDMjhCQzI5MTFDQzNBNTI1NDMyQkMzPjw0RUYwMEI5Q0NEMTg5Qzc0NjVDQ0ZEM0IzRTY4NUVGNT5dPj4Kc3RhcnR4cmVmCjYzMQolJUVPRgo="
        )
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(pdf_bytes)
            path = tmp.name

        try:
            parsed = parse_resume_structured(path, fast=True)
            self.assertTrue(parsed.get("metadata", {}).get("fast"))
            self.assertIn("Software Engineer", parsed.get("text", ""))
            self.assertEqual(parsed.get("metadata", {}).get("ocr_pages"), 0)
        finally:
            import os
            os.remove(path)

    def test_multicolumn_docx_and_links_are_preserved(self):
        from docx import Document

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "synthetic.docx"
            document = Document()
            document.add_paragraph("Synthetic Candidate")
            table = document.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "EXPERIENCE\nSoftware Engineer"
            table.cell(0, 1).text = "SKILLS\nPython, SQLite"
            document.add_paragraph("github.com/synthetic-candidate")
            document.save(path)

            parsed = parse_resume_structured(str(path), fast=True)

        self.assertIn("Software Engineer", parsed["text"])
        self.assertIn("Python, SQLite", parsed["text"])
        self.assertIn("https://github.com/synthetic-candidate", parsed["links"])

    def test_scanned_pdf_uses_optional_ocr_result_when_available(self):
        import pypdfium2 as pdfium

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "scan.pdf"
            document = pdfium.PdfDocument.new()
            page = document.new_page(595, 842)
            page.close()
            document.save(path)
            document.close()
            with patch("resume_parser._ocr_pdf_pages", return_value=["Synthetic Candidate Python Engineer"]):
                parsed = parse_resume_structured(str(path), fast=False)

        self.assertEqual(parsed["parser"], "pypdfium2+paddleocr")
        self.assertEqual(parsed["metadata"]["ocr_pages"], 1)
        self.assertIn("Synthetic Candidate", parsed["text"])

    def test_oversized_pdf_page_is_rejected_before_rendering(self):
        import pypdfium2 as pdfium

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "oversized.pdf"
            document = pdfium.PdfDocument.new()
            page = document.new_page(10_000, 10_000)
            page.close()
            document.save(path)
            document.close()

            with self.assertRaisesRegex(ValueError, "page is too large"):
                parse_resume_structured(str(path), fast=True)

    def test_malformed_pdf_fails_with_a_safe_message(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "malformed.pdf"
            path.write_bytes(b"%PDF-1.7\nnot-a-real-document")

            with self.assertRaisesRegex(RuntimeError, "Upload a valid PDF or DOCX"):
                parse_resume_structured(str(path), fast=True)

    def test_password_encrypted_pdf_fails_without_exposing_parser_details(self):
        encrypted_pdf = base64.b64decode(
            "JVBERi0xLjcKJeLjz9MKMSAwIG9iago8PAovUHJvZHVjZXIgPGZkNTBiNGFmNjY+Cj4+CmVuZGJqCjIgMCBvYmoKPDwKL1R5cGUgL1BhZ2VzCi9Db3VudCAxCi9LaWRzIFsgNCAwIFIgXQo+PgplbmRvYmoKMyAwIG9iago8PAovVHlwZSAvQ2F0YWxvZwovUGFnZXMgMiAwIFIKPj4KZW5kb2JqCjQgMCBvYmoKPDwKL1R5cGUgL1BhZ2UKL01lZGlhQm94IFsgMCAwIDU5NSA4NDIgXQovUm90YXRlIDAKL1Jlc291cmNlcyA1IDAgUgovQ29udGVudHMgWyA3IDAgUiBdCi9QYXJlbnQgMiAwIFIKPj4KZW5kb2JqCjUgMCBvYmoKPDwKL0ZvbnQgPDwKL2hlbHYgNiAwIFIKPj4KPj4KZW5kb2JqCjYgMCBvYmoKPDwKL1R5cGUgL0ZvbnQKL1N1YnR5cGUgL1R5cGUxCi9CYXNlRm9udCAvSGVsdmV0aWNhCi9FbmNvZGluZyAvV2luQW5zaUVuY29kaW5nCj4+CmVuZG9iago3IDAgb2JqCjw8Ci9GaWx0ZXIgL0ZsYXRlRGVjb2RlCi9MZW5ndGggMTUzCj4+CnN0cmVhbQpBVEJ7U1AyP1pON0tGck9GTFlPby9zUXB6U3phR3FlUHgycUo5Y0w1N29JLzIyc2xaMU5Rd1UyajVxbGF0c1RmNm54M3BqeDFBRUp3ZW9JUWpnMzh4NjJkMlJwL29mYmY3dGdNOVpGd3k2TFVQVTlHWmw0MGFPcHpoM2RLbGIrNlJsUFpuVWFISDM4UVd5TisvWlpFdkd1dWh4WjBtV0NReEl4clVmNnRyNDFBckFiaUEyWlkyWUtIVDlLeEd0OEJ3eDJkbmZLWDhrQW4KZW5kc3RyZWFtCmVuZG9iago4IDAgb2JqCjw8Ci9WIDIKL1IgMwovTGVuZ3RoIDEyOAovUCA0Mjk0OTY3MjkyCi9GaWx0ZXIgL1N0YW5kYXJkCi9PIDwwZTUyMjkyNWEzZTRlODc0YzNjZmFjYmVmNTExYTczYWM0ZWMyYmQ4NjVkY2QzZDQ2Mjc2MTQ5MTdhYmZkN2U0PgovVSA8ZWJlYjI0NGRiNjQ4MGMzZjM1MjNkYjM0Yzk2MWU2MjAyOGJmNGU1ZTRlNzU4YTQxNjQwMDRlNTZmZmZhMDEwOD4KPj4KZW5kb2JqCnhyZWYKMCA5CjAwMDAwMDAwMDAgNjU1MzUgZiAKMDAwMDAwMDAxNSAwMDAwMCBuIAowMDAwMDAwMDU5IDAwMDAwIG4gCjAwMDAwMDAxMTggMDAwMDAgbiAKMDAwMDAwMDE2NyAwMDAwMCBuIAowMDAwMDAwMjg3IDAwMDAwIG4gCjAwMDAwMDAzMzIgMDAwMDAgbiAKMDAwMDAwMDQyOSAwMDAwMCBuIAowMDAwMDAwNjU0IDAwMDAwIG4gCnRyYWlsZXIKPDwKL1NpemUgOQovUm9vdCAzIDAgUgovSW5mbyAxIDAgUgovSUQgWyA8MzkzNjM0Mzk2NjMyNjMzNDY2NjEzMTMzMzU2MTMwMzAzNTYzMzUzMDM0NjQzMTY1MzQ2NTYyNjQzODMzMzg2NT4gPDM5MzYzNDM5NjYzMjYzMzQ2NjYxMzEzMzM1NjEzMDMwMzU2MzM1MzAzNDY0MzE2NTM0NjU2MjY0MzgzMzM4NjU+IF0KL0VuY3J5cHQgOCAwIFIKPj4Kc3RhcnR4cmVmCjg2OQolJUVPRgo="
        )
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "encrypted.pdf"
            path.write_bytes(encrypted_pdf)

            with self.assertRaisesRegex(RuntimeError, "Upload a valid PDF or DOCX"):
                parse_resume_structured(str(path), fast=True)


if __name__ == "__main__":
    unittest.main()
