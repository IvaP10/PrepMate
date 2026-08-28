# ============================================================================
# MODULE: resume_parser.py
# PURPOSE: Extract text + links + light metadata from PDF/DOCX resumes
#          (PDFium for selectable PDFs, PaddleOCR fallback for scans, python-docx
#          for Word documents).
# STRUCTURE:
#   - ResumeParseResult dataclass (lines 23-32)
#   - URL_PATTERN regex (line 35-38)
#   - _normalize_text helper + per-format parsers (later in file)
#   - parse_resume_structured(file_path) entry point
# ENDPOINTS: none
# DEPENDS ON: pypdfium2, paddleocr, python-docx (lazy imports)
# CONSUMED BY: pre_interview.py
# DATA TABLES: none (returns a dict; pre_interview.py persists into UserInfo.resume_json)
# ============================================================================

from __future__ import annotations

import logging
import io
import os
import re
import tempfile
import zipfile
from dataclasses import dataclass, asdict
from pathlib import PurePosixPath
from typing import Any, Dict, Iterable, List

logger = logging.getLogger("resume_parser")

MAX_RESUME_PDF_PAGES = 20
MAX_RESUME_PDF_PAGE_PIXELS = 20_000_000
MAX_RESUME_PDF_TOTAL_PIXELS = 80_000_000
MAX_DOCX_ARCHIVE_ENTRIES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 20 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 100


def validate_resume_bytes(content: bytes, extension: str) -> None:
    """Validate the container before handing bytes to native document parsers."""
    ext = str(extension or "").lower()
    if ext == ".pdf":
        if not content.startswith(b"%PDF-"):
            raise ValueError("The uploaded file is not a valid PDF.")
        return
    if ext != ".docx" or not content.startswith(b"PK"):
        raise ValueError("The uploaded file is not a valid DOCX archive.")

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            names = {member.filename for member in members}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise ValueError("The uploaded archive is not a DOCX document.")
            if len(members) > MAX_DOCX_ARCHIVE_ENTRIES:
                raise ValueError("The DOCX archive contains too many files.")
            total_uncompressed = sum(max(0, member.file_size) for member in members)
            if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ValueError("The DOCX archive expands beyond the allowed size.")
            for member in members:
                member_name = member.filename.replace("\\", "/")
                member_path = PurePosixPath(member_name)
                if member_path.is_absolute() or ".." in member_path.parts:
                    raise ValueError("The DOCX archive contains an unsafe file path.")
                if member.file_size <= 0:
                    continue
                compressed = max(1, member.compress_size)
                if member.file_size / compressed > MAX_DOCX_COMPRESSION_RATIO:
                    raise ValueError("The DOCX archive has an unsafe compression ratio.")
                if member_name.lower().endswith((".xml", ".rels")):
                    xml_content = archive.read(member).upper()
                    if b"<!DOCTYPE" in xml_content or b"<!ENTITY" in xml_content:
                        raise ValueError("The DOCX archive contains unsafe XML declarations.")
    except zipfile.BadZipFile as exc:
        raise ValueError("The uploaded file is not a valid DOCX archive.") from exc


@dataclass
class ResumeParseResult:
    text: str
    parser: str 
    links: List[str]
    metadata: Dict[str, Any]

    def to_dict(self) -> Dict[str, Any]:
        return asdict(self)


URL_PATTERN = re.compile(
    r"(?:https?://)?(?:www\.)?(?:linkedin\.com/in/[A-Za-z0-9_.-]+|github\.com/[A-Za-z0-9_.-]+|[A-Za-z0-9_.-]+\.(?:dev|io|me|com|net|org)(?:/[^\s),;]*)?)",
    re.I,
)


def _normalize_text(text: str) -> str:
    text = text or ""
    text = text.replace("&amp;", "&").replace("&lt;", "<").replace("&gt;", ">")
    text = re.sub(r"\b([A-Z])\s+(?=[A-Z]\b)", r"\1", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n", text)
    return text.strip()


def _extract_links(text: str) -> List[str]:
    links: List[str] = []
    for match in URL_PATTERN.finditer(text or ""):
        url = match.group(0).strip().rstrip(".,;)")
        if not url:
            continue
        if "." in url and not url.startswith("http"):
            url = f"https://{url}"
        links.append(url)
    return list(dict.fromkeys(links))


def _page_text_pdfium(page) -> str:
    text_page = page.get_textpage()
    try:
        return text_page.get_text_range(force_this=True) or ""
    finally:
        text_page.close()


def _parse_pdf_with_pdfium(file_path: str, *, allow_ocr: bool = True) -> ResumeParseResult:
    import pypdfium2 as pdfium

    doc = pdfium.PdfDocument(file_path)
    try:
        if len(doc) <= 0 or len(doc) > MAX_RESUME_PDF_PAGES:
            raise ValueError(f"PDF resumes must contain between 1 and {MAX_RESUME_PDF_PAGES} pages.")
        total_pixels = 0
        for page_index in range(len(doc)):
            page = doc[page_index]
            width, height = page.get_size()
            page_pixels = int(max(0, width * 2) * max(0, height * 2))
            if page_pixels > MAX_RESUME_PDF_PAGE_PIXELS:
                raise ValueError("A PDF page is too large to process safely.")
            total_pixels += page_pixels
            page.close()
        if total_pixels > MAX_RESUME_PDF_TOTAL_PIXELS:
            raise ValueError("The PDF exceeds the total page pixel budget.")
        text_pages = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            try:
                text_pages.append(_page_text_pdfium(page))
            finally:
                page.close()
        ocr_pages: List[str] = []
        normalized_preview = _normalize_text("\n".join(text_pages))
        low_text_pages = [
            index
            for index, text in enumerate(text_pages)
            if len((text or "").strip()) < 80
        ]
        if allow_ocr and (low_text_pages or len(normalized_preview) < 250):
            ocr_pages = _ocr_pdf_pages(doc, low_text_pages or list(range(len(doc))))
        pages = text_pages + ocr_pages
        page_count = len(doc)
    finally:
        doc.close()
    text = _normalize_text("\n".join(pages))
    if not text:
        raise ValueError("PDF text/OCR extraction returned empty text")
    parser = "pypdfium2+paddleocr" if ocr_pages else "pypdfium2"
    return ResumeParseResult(
        text=text,
        parser=parser,
        links=_extract_links(text),
        metadata={
            "pages": page_count,
            "ocr_pages": len(ocr_pages),
            "format": ".pdf",
        },
    )


_paddle_ocr = None
_paddle_loaded = False


def _get_paddle_ocr():
    global _paddle_ocr, _paddle_loaded
    if _paddle_loaded:
        return _paddle_ocr
    _paddle_loaded = True
    try:
        from paddleocr import PaddleOCR  # type: ignore[import-not-found]

        try:
            _paddle_ocr = PaddleOCR(
                lang="en",
                use_doc_orientation_classify=False,
                use_doc_unwarping=False,
                use_textline_orientation=True,
            )
        except TypeError:
            _paddle_ocr = PaddleOCR(lang="en", use_angle_cls=True)
        logger.info("PaddleOCR resume parser loaded")
    except Exception:
        logger.error("PaddleOCR unavailable")
        _paddle_ocr = None
    return _paddle_ocr


def _ocr_pdf_pages(doc, page_indexes: Iterable[int]) -> List[str]:
    ocr = _get_paddle_ocr()
    if not ocr:
        return []

    extracted: List[str] = []
    for page_index in page_indexes:
        tmp_path = ""
        page = None
        try:
            page = doc[page_index]
            pix = page.render(scale=2).to_pil()
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_path = tmp.name
            pix.save(tmp_path, format="PNG")
            result = _run_paddle_ocr(ocr, tmp_path)
            page_text = _normalize_text("\n".join(_flatten_ocr_result(result)))
            if page_text:
                extracted.append(page_text)
        except Exception:
            logger.debug("PaddleOCR page %s failed", page_index)
        finally:
            if page is not None:
                try:
                    page.close()
                except Exception:
                    pass
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass
    return extracted


def _run_paddle_ocr(ocr, image_path: str):
    if hasattr(ocr, "predict"):
        return ocr.predict(input=image_path)
    return ocr.ocr(image_path, cls=True)


def _flatten_ocr_result(value: Any) -> List[str]:
    texts: List[str] = []

    def visit(obj: Any) -> None:
        if obj is None:
            return
        if isinstance(obj, str):
            if re.search(r"[A-Za-z]", obj) and len(obj.strip()) > 1:
                texts.append(obj.strip())
            return
        if isinstance(obj, dict):
            for key in ("rec_text", "text", "transcription"):
                if isinstance(obj.get(key), str):
                    visit(obj[key])
            for key in ("rec_texts", "texts", "data", "res", "result"):
                if key in obj:
                    visit(obj[key])
            return
        if isinstance(obj, (list, tuple)):
            if len(obj) >= 2 and isinstance(obj[1], (list, tuple)) and obj[1] and isinstance(obj[1][0], str):
                visit(obj[1][0])
                return
            for item in obj:
                visit(item)
            return
        for attr in ("rec_text", "rec_texts", "text", "json"):
            try:
                if hasattr(obj, attr):
                    visit(getattr(obj, attr))
            except Exception:
                pass

    visit(value)
    return list(dict.fromkeys(texts))


def _parse_docx_with_python_docx(file_path: str) -> ResumeParseResult:
    from docx import Document  # type: ignore[import-not-found]

    with open(file_path, "rb") as resume_file:
        validate_resume_bytes(resume_file.read(), ".docx")
    doc = Document(file_path)
    parts = [para.text for para in doc.paragraphs if para.text and para.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = _normalize_text("\n".join(parts))
    if not text:
        raise ValueError("DOCX extraction returned empty text")
    return ResumeParseResult(
        text=text,
        parser="python-docx",
        links=_extract_links(text),
        metadata={"format": ".docx"},
    )


def parse_resume_structured(file_path: str, *, fast: bool = False) -> Dict[str, Any]:
    if not os.path.exists(file_path):
        raise FileNotFoundError("Resume file not found")

    ext = os.path.splitext(file_path)[1].lower()
    errors: List[str] = []

    with open(file_path, "rb") as resume_file:
        validate_resume_bytes(resume_file.read(), ext)

    try:
        if ext == ".pdf":
            result = _parse_pdf_with_pdfium(file_path, allow_ocr=not fast)
        elif ext == ".docx":
            result = _parse_docx_with_python_docx(file_path)
        else:
            raise ValueError("No parser for this file type")
        parsed = result.to_dict()
        parsed["metadata"]["fast"] = fast
        parsed["metadata"]["errors"] = errors
        return parsed
    except ValueError:
        raise
    except Exception as exc:
        errors.append(f"parser: {type(exc).__name__}")
        logger.error("All resume parsers failed")
        raise RuntimeError(
            "Failed to read resume. Upload a valid PDF or DOCX, or convert legacy .doc files to DOCX."
        ) from exc


def parse_resume(file_path: str) -> str:
    return parse_resume_structured(file_path)["text"]
