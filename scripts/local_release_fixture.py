#!/usr/bin/env python3
"""Seed or remove a disposable, authenticated local browser-release fixture."""

from __future__ import annotations

import argparse
from datetime import datetime, timedelta, timezone
import json
import os
from pathlib import Path
import secrets
import sys
import tempfile
import uuid

import httpx
from docx import Document
from dotenv import load_dotenv


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT))
load_dotenv(ROOT / "key.env")

from config import settings  # noqa: E402
from database import get_db, init_connection_pool, close_connection_pool  # noqa: E402
from learning_engine import _ensure_active_improvement_mission  # noqa: E402
from security_utils import encrypt_data  # noqa: E402


def _require_local(api_base_url: str) -> None:
    if settings.ENVIRONMENT == "production":
        raise SystemExit("Local release fixtures are forbidden in production")
    host = httpx.URL(api_base_url).host
    if host not in {"localhost", "127.0.0.1", "::1"}:
        raise SystemExit("Fixture API URL must resolve to localhost")


def _request(client: httpx.Client, method: str, path: str, **kwargs) -> httpx.Response:
    response = client.request(method, path, **kwargs)
    if response.status_code >= 400:
        raise RuntimeError(f"{method} {path} failed with HTTP {response.status_code}: {response.text[:300]}")
    return response


def seed(api_base_url: str, output: Path) -> dict:
    _require_local(api_base_url)
    suffix = uuid.uuid4().hex[:10]
    email = f"release-e2e-{suffix}@example.test"
    password = f"E2e!{secrets.token_urlsafe(18)}aA1"
    name = "Release Verification Candidate"

    with httpx.Client(base_url=api_base_url, follow_redirects=False, timeout=20) as client:
        _request(client, "POST", "/api/auth/signup", json={"name": name, "email": email, "password": password})
        with get_db() as connection:
            cursor = connection.cursor()
            cursor.execute("SELECT verification_token FROM Login WHERE email = %s", (email,))
            verification_token = (cursor.fetchone() or [None])[0]
            cursor.close()
        if not verification_token:
            raise RuntimeError("Signup did not persist an email verification token")
        verification = _request(client, "GET", f"/api/auth/verify-email?token={verification_token}")
        if verification.status_code not in {302, 303, 307, 308}:
            raise RuntimeError("Email verification did not return the expected redirect")
        login = _request(client, "POST", "/api/auth/login", json={"email": email, "password": password})
        login_payload = login.json()
        user_id = str(login_payload["user_id"])
        csrf_token = client.cookies.get("interai_csrf")
        if not csrf_token:
            raise RuntimeError("Login did not issue the CSRF cookie required for mutations")
        _request(
            client,
            "POST",
            "/api/pre-interview/confirm-profile",
            headers={"X-CSRF-Token": csrf_token},
            json={
                "job_id": None,
                "profile": {
                    "name": name,
                    "email": email,
                    "phone": "+91 90000 00000",
                    "skills": ["Python", "FastAPI", "PostgreSQL", "System Design"],
                    "target_role": "Backend Engineer",
                    "summary": "Backend engineer building reliable APIs and evidence-based learning systems.",
                },
            },
        )
        with tempfile.TemporaryDirectory(prefix="interai-release-") as temp_dir:
            resume_path = Path(temp_dir) / "release-verification-resume.docx"
            document = Document()
            document.add_heading(name, 0)
            document.add_paragraph(email)
            document.add_heading("Summary", level=1)
            document.add_paragraph("Backend engineer building reliable APIs and evidence-based learning systems.")
            document.add_heading("Skills", level=1)
            document.add_paragraph("Python, FastAPI, PostgreSQL, System Design")
            document.add_heading("Experience", level=1)
            document.add_paragraph(
                "Backend Engineer at Example Systems. Reduced API p95 latency by 38 percent "
                "using idempotent jobs and database indexing."
            )
            document.save(resume_path)
            with resume_path.open("rb") as resume_file:
                resume_upload = _request(
                    client,
                    "POST",
                    "/api/pre-interview/upload-resume",
                    headers={"X-CSRF-Token": csrf_token},
                    files={
                        "file": (
                            resume_path.name,
                            resume_file,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        ),
                    },
                ).json()
        resume_id = str(resume_upload["resume"]["resume_id"])
        job_profile = _request(
            client,
            "POST",
            "/api/workspace/job-profiles",
            headers={"X-CSRF-Token": csrf_token},
            json={
                "role": "Backend Engineer",
                "company": "Example Systems",
                "tech_stack": ["Python", "FastAPI", "PostgreSQL"],
                "job_description": (
                    "Build reliable backend APIs, design distributed systems, own database performance, "
                    "and communicate engineering trade-offs."
                ),
                "experience_level": "mid",
                "requirements": ["API design", "system design", "database performance"],
            },
        ).json()
        job_profile_id = int(job_profile["profile_id"])

    now = datetime.now(timezone.utc).replace(tzinfo=None)
    interview_id = str(uuid.uuid4())
    technical_id = str(uuid.uuid4())
    report_id = str(uuid.uuid4())
    round_id = str(uuid.uuid4())
    round_spec_id = str(uuid.uuid4())
    question_id = str(uuid.uuid4())
    response_id = str(uuid.uuid4())
    assessment_id = str(uuid.uuid4())

    report_payload = {
        "version": "report-v2",
        "report_state": "ready",
        "summary": "You communicated the architecture clearly and supported the main decision with measurable evidence.",
        "readiness_label": "Ready with targeted practice",
        "overall_score": 82,
        "skill_scores": {"communication": 84, "problem_solving": 80, "technical_depth": 81},
        "topic_breakdown": [{"topic": "Project ownership", "score": 82, "turns": 1}],
        "behavioral_metrics": {"question_count": 1, "average_response_time_seconds": 54},
        "student_summary": {
            "headline": "Strong evidence-backed explanation",
            "interviewer_signal": "You owned the decision and quantified its effect.",
            "next_step": "Practise stating the trade-off earlier.",
        },
        "strengths": ["Clear ownership", "Measurable result"],
        "improvements": [{"title": "Lead with the trade-off", "detail": "State the rejected alternative before implementation detail."}],
        "practice_plan": [{"day": "Next practice", "task": "Repeat the answer in 90 seconds with the trade-off in the first half."}],
        "findings": [{
            "finding_key": "project-ownership",
            "what_happened": "The answer connected a concrete engineering decision to a measurable latency reduction.",
            "why_matters": "This gives interviewers direct evidence of ownership and impact.",
            "evidence_ids": [response_id],
            "confidence": 0.91,
            "recommended_action": "Keep the same evidence and state the trade-off earlier.",
            "measurement": "Direct answer and trade-off within the first 30 seconds.",
        }],
        "per_turn_feedback": [{
            "response_id": response_id,
            "question_id": question_id,
            "question": "Tell me about a backend decision you owned and the impact it produced.",
            "question_type": "main",
            "is_followup": False,
            "topic": "Project ownership",
            "response": "I introduced idempotent job processing and a targeted database index, reducing p95 latency by 38 percent while preventing duplicate work.",
            "score": 82,
            "feedback": "Strong ownership and measurable evidence. State the alternative and trade-off sooner.",
            "time_taken": 54,
            "coaching_hint": None,
            "answer_quality_flags": [],
            "evidence_quotes": ["reducing p95 latency by 38 percent"],
        }],
    }

    with get_db() as connection:
        cursor = connection.cursor()
        try:
            mission_id = _ensure_active_improvement_mission(
                cursor,
                user_id,
                [{
                    "skill_key": "communication:project_ownership",
                    "label": "Project ownership",
                    "category": "communication",
                    "mastery_score": 54,
                    "confidence_score": 76,
                    "evidence_count": 3,
                    "why_it_matters": "Your evidence was useful, but the decision and trade-off arrived too late.",
                }],
                [],
                [],
                mode="mock",
                source_interview_id=None,
                source_analysis_id=None,
            )
            cursor.execute(
                """
                SELECT roadmap_node_id, exercise_id
                FROM ImprovementRoadmapNodes
                WHERE mission_id = %s AND user_id = %s AND availability_status = 'current'
                """,
                (mission_id, user_id),
            )
            roadmap_node_id, exercise_id = cursor.fetchone()

            common_questions = {
                "duration_minutes": 45,
                "battlegrounds": [{
                    "id": 1,
                    "section_id": "project-ownership",
                    "label": "Project ownership",
                    "kind": "resume",
                    "opening_question": "Tell me about a backend decision you owned and the impact it produced.",
                    "importance": "critical",
                    "estimated_difficulty": "matched",
                    "min_turns": 1,
                    "max_turns": 2,
                    "current_turns": 0,
                    "time_budget_seconds": 600,
                    "expected_points": ["personal contribution", "measurable result", "trade-off"],
                    "taxonomy_keys": ["behavioral:ownership"],
                    "rubric": {"version": "browser-e2e-v1", "unknown_dimensions_are_null": True},
                }],
            }
            cursor.executemany(
                """
                INSERT INTO Interviews (
                    interview_id, user_id, interview_mode, interview_type, job_title,
                    strictness_level, status, session_id, persona_data, questions_data,
                    settings, started_at, deadline_at, attempt_status, analysis_status,
                    integrity_status, lifecycle_revision
                ) VALUES (%s, %s, %s, %s, %s, 'medium', %s, %s, %s, %s, %s,
                          %s, %s, 'active', 'not_requested', 'clean', 1)
                """,
                [
                    (
                        interview_id, user_id, "mock", "Mock Interview", "Backend Engineer",
                        "in_progress", f"session-{suffix}", json.dumps({"name": "Ava", "personality": "Professional"}),
                        json.dumps(common_questions), json.dumps({
                            "profile_type": "mid_tier", "input_mode": "voice", "camera_mode": "required",
                            "duration_minutes": 45, "job_title": "Backend Engineer",
                        }), now, now + timedelta(minutes=45),
                    ),
                    (
                        technical_id, user_id, "mock", "technical", "Backend Engineer",
                        "in_progress", f"technical-session-{suffix}", json.dumps({"name": "Ava", "personality": "Rigorous"}),
                        json.dumps({"duration_minutes": 50, "battlegrounds": []}), json.dumps({
                            "profile_type": "mid_tier", "programming_language": "python",
                            "technical_rounds": ["coding"], "duration_minutes": 50,
                        }), now, now + timedelta(minutes=50),
                    ),
                ],
            )
            cursor.execute(
                """
                INSERT INTO TechnicalInterviewRounds (
                    round_id, interview_id, user_id, round_type, language, prompt,
                    starter_code, whiteboard_json, status, metadata, round_spec_id,
                    problem_id, round_number, round_spec, duration_seconds, deadline_at,
                    mode, max_submissions, problem_version, workflow_state, started_at
                ) VALUES (%s, %s, %s, 'coding', 'python', %s, %s, '{}'::jsonb, 'active',
                          %s, %s, %s, 1, %s, 1800, %s, 'mock', 1, 1, '{}'::jsonb, %s)
                """,
                (
                    round_id, technical_id, user_id,
                    "Implement a function that returns the first non-repeating character in a string.",
                    "def first_unique(text: str):\n    # Write your solution here\n    pass\n",
                    json.dumps({"title": "First unique character", "visible_tests": [{"input": "swiss", "expected": "w"}]}),
                    round_spec_id, None, json.dumps({
                        "title": "First unique character", "prompt": "Return the first character that appears once.",
                        "visible_tests": [{"input": "swiss", "expected": "w"}],
                    }), now + timedelta(minutes=30), now,
                ),
            )
            cursor.execute(
                """
                INSERT INTO Interviews (
                    interview_id, user_id, interview_mode, interview_type, job_title,
                    strictness_level, status, session_id, persona_data, questions_data,
                    settings, overall_score, feedback_summary, completed_at, report_json,
                    report_json_encrypted, duration_seconds, started_at, deadline_at,
                    attempt_status, analysis_status, integrity_status, lifecycle_revision,
                    completion_kind, evidence_sealed_at, evidence_hash
                ) VALUES (%s, %s, 'mock', 'Mock Interview', 'Backend Engineer', 'medium',
                          'completed', %s, '{}'::jsonb, '{}'::jsonb, '{}'::jsonb, 82, %s,
                          %s, %s, %s, 1740, %s, %s, 'completed', 'ready', 'clean', 3,
                          'natural', %s, %s)
                """,
                (
                    report_id, user_id, f"report-session-{suffix}", report_payload["summary"], now,
                    json.dumps({"encrypted": True}), encrypt_data(json.dumps(report_payload)).encode("utf-8"),
                    now - timedelta(minutes=29), now, None, None,
                ),
            )
            cursor.execute(
                """
                INSERT INTO InterviewQuestions (
                    question_id, interview_id, question_text, question_order, question_type,
                    difficulty_level, is_followup, topic_label, profile_type, rubric_version,
                    source, quality_score, generation_metadata, taxonomy_keys, expected_points,
                    rubric_json, selection_reason, blueprint_section_id, provenance,
                    question_spec_id, max_followups, time_budget_seconds, claim_ids,
                    expected_point_ids
                ) VALUES (%s, %s, %s, 1, 'main', 'medium', FALSE, 'Project ownership',
                          'mid_tier', 'browser-e2e-v1', 'fixture', 1, '{}'::jsonb,
                          '["behavioral:ownership"]'::jsonb, '[]'::jsonb, '{}'::jsonb,
                          'browser_release_verification', 'project-ownership', '{}'::jsonb,
                          %s, 2, 600, '[]'::jsonb, '[]'::jsonb)
                """,
                (question_id, report_id, report_payload["per_turn_feedback"][0]["question"], f"question-spec-{suffix}"),
            )
            answer = report_payload["per_turn_feedback"][0]["response"]
            assessment = {
                "version": "evaluation-v1", "overall_score": 82, "insufficient_evidence": False,
                "feedback": report_payload["per_turn_feedback"][0]["feedback"],
            }
            cursor.execute(
                """
                INSERT INTO InterviewResponses (
                    response_id, interview_id, question_id, user_response, response_time_seconds,
                    ai_feedback, score, nonverbal_metrics, evaluation_json,
                    answer_quality_flags, evidence_quotes, idempotency_key, evidence_hash,
                    answer_text_encrypted, raw_answer_hash, input_mode, timing_json
                ) VALUES (%s, %s, %s, '[encrypted]', 54, %s, 82, '{}'::jsonb, %s,
                          '[]'::jsonb, %s, %s, %s, %s, %s, 'voice', '{}'::jsonb)
                """,
                (
                    response_id, report_id, question_id, assessment["feedback"], json.dumps(assessment),
                    json.dumps(report_payload["per_turn_feedback"][0]["evidence_quotes"]),
                    f"response-idempotency-{suffix}", f"response-evidence-{suffix}",
                    encrypt_data(answer).encode("utf-8"), str(uuid.uuid4()),
                ),
            )
            cursor.execute(
                """
                INSERT INTO ResponseAssessments (
                    assessment_id, response_id, interview_id, evaluator_version,
                    evidence_hash, overall_score, assessment_json
                ) VALUES (%s, %s, %s, 'evaluation-v1', %s, 82, %s)
                """,
                (assessment_id, response_id, report_id, f"assessment-evidence-{suffix}", json.dumps(assessment)),
            )
            cursor.execute(
                "UPDATE Interviews SET evidence_sealed_at = %s, evidence_hash = %s WHERE interview_id = %s",
                (now, f"evidence-{suffix}", report_id),
            )
            connection.commit()
        except Exception:
            connection.rollback()
            raise
        finally:
            cursor.close()

    fixture = {
        "email": email,
        "password": password,
        "user_id": user_id,
        "interview_id": interview_id,
        "technical_id": technical_id,
        "report_id": report_id,
        "mission_id": str(mission_id),
        "roadmap_node_id": str(roadmap_node_id),
        "exercise_id": str(exercise_id),
        "resume_id": resume_id,
        "job_profile_id": job_profile_id,
        "api_base_url": api_base_url,
    }
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(fixture, indent=2) + "\n", encoding="utf-8")
    return fixture


def cleanup(api_base_url: str, fixture_path: Path) -> None:
    _require_local(api_base_url)
    fixture = json.loads(fixture_path.read_text(encoding="utf-8"))
    with httpx.Client(base_url=api_base_url, timeout=20) as client:
        login = _request(client, "POST", "/api/auth/login", json={
            "email": fixture["email"], "password": fixture["password"],
        })
        csrf_token = client.cookies.get("interai_csrf")
        if not csrf_token:
            raise RuntimeError("Login did not issue the CSRF cookie required for account deletion")
        _request(
            client,
            "DELETE",
            "/api/auth/delete-account",
            headers={"X-CSRF-Token": csrf_token},
            json={"password": fixture["password"]},
        )
    fixture_path.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("command", choices=("seed", "cleanup"))
    parser.add_argument("--api-base-url", default="http://localhost:8000")
    parser.add_argument(
        "--output",
        type=Path,
        default=ROOT / "Frontend/e2e/.generated/local-release-fixture.json",
    )
    args = parser.parse_args()
    init_connection_pool()
    try:
        if args.command == "seed":
            print(json.dumps(seed(args.api_base_url.rstrip("/"), args.output), sort_keys=True))
        else:
            cleanup(args.api_base_url.rstrip("/"), args.output)
            print(json.dumps({"cleaned": True, "fixture": str(args.output)}, sort_keys=True))
        return 0
    finally:
        close_connection_pool()


if __name__ == "__main__":
    raise SystemExit(main())
